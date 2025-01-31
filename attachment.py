import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import io
from PIL import Image
import zipfile
import os
import shutil
import tempfile

# Define functions
def set_font(paragraph, font_name, font_size):
  for run in paragraph.runs:
      run.font.name = font_name
      r = run._element
      rFonts = r.rPr.rFonts
      rFonts.set(qn('w:eastAsia'), font_name)
      run.font.size = Pt(font_size)

def set_document_font(doc, font_name='Times New Roman', font_size=12):
  for paragraph in doc.paragraphs:
      set_font(paragraph, font_name, font_size)
  for table in doc.tables:
      for row in table.rows:
          for cell in row.cells:
              for paragraph in cell.paragraphs:
                  set_font(paragraph, font_name, font_size)

def append_pdf_to_docx(pdf_paths, docx_path):
  try:
      doc = Document(docx_path)
  except:
      doc = Document()

  for pdf_path in pdf_paths:
      pdf_document = fitz.open(pdf_path)

      for page_num in range(len(pdf_document)):
          page = pdf_document.load_page(page_num)
          pix = page.get_pixmap()
          img_data = pix.tobytes()
          image = Image.open(io.BytesIO(img_data))
          image_stream = io.BytesIO()
          image.save(image_stream, format='PNG')
          doc.add_picture(image_stream, width=Inches(6))

          if page_num < len(pdf_document) - 1:
              doc.add_page_break()

  set_document_font(doc)
  doc.save(docx_path)

def append_docx_to_docx(docx_paths, docx_path):
  try:
      doc = Document(docx_path)
  except:
      doc = Document()

  for path in docx_paths:
      sub_doc = Document(path)
      for element in sub_doc.element.body:
          doc.element.body.append(element)

  set_document_font(doc)
  doc.save(docx_path)

def process_zip(zip_path, file_paths, output_zip_path, file_type):
  extract_dir = tempfile.mkdtemp()

  with zipfile.ZipFile(zip_path, 'r') as zip_ref:
      zip_ref.extractall(extract_dir)

  for root, _, files in os.walk(extract_dir):
      for file in files:
          if file.endswith('.docx'):
              docx_path = os.path.join(root, file)
              if file_type == 'pdf':
                  append_pdf_to_docx(file_paths, docx_path)
              elif file_type == 'docx':
                  append_docx_to_docx(file_paths, docx_path)

  with zipfile.ZipFile(output_zip_path, 'w') as zip_ref:
      for root, _, files in os.walk(extract_dir):
          for file in files:
              full_path = os.path.join(root, file)
              zip_ref.write(full_path, os.path.relpath(full_path, extract_dir))

  shutil.rmtree(extract_dir)

def process_appended_zip(zip_path, output_zip_path):
  extract_dir = tempfile.mkdtemp()

  with zipfile.ZipFile(zip_path, 'r') as zip_ref:
      zip_ref.extractall(extract_dir)

  docx_files = []
  pdf_files = []

  for root, _, files in os.walk(extract_dir):
      for file in files:
          full_path = os.path.join(root, file)
          if file.endswith('.docx'):
              docx_files.append(full_path)
          elif file.endswith('.pdf'):
              pdf_files.append(full_path)

  if docx_files:
      process_zip(zip_path, docx_files, output_zip_path, 'docx')
  if pdf_files:
      process_zip(zip_path, pdf_files, output_zip_path, 'pdf')

  shutil.rmtree(extract_dir)

# Streamlit app
st.title("Append PDFs or Word Documents to Word Documents")

# Add a text box with a link to MEA Attachment A
st.markdown("### MEA Attachment A available here: [MEA Attachment A](https://energy.maryland.gov/Pages/all-incentives.aspx)")

# Upload zip file
zip_file = st.file_uploader("Upload ZIP file containing Word or PDF documents", type=["zip"])

# Choose file type to append
file_type = st.selectbox("Select file type to append", ["PDF", "Word Document", "Both PDF and Word Document"])

# Upload multiple files
file_types = {"PDF": "pdf", "Word Document": "docx", "Both PDF and Word Document": ["pdf", "docx"]}
file_extension = file_types[file_type]
if isinstance(file_extension, list):
  files = st.file_uploader(f"Upload PDF and Word Document files to append", type=file_extension, accept_multiple_files=True)
else:
  files = st.file_uploader(f"Upload {file_type} files to append", type=[file_extension], accept_multiple_files=True)

# Option to upload a zip file containing Word or PDF documents to append
append_zip_file = st.file_uploader("Upload ZIP file containing documents to append", type=["zip"])

if st.button("Process"):
  if zip_file and (files or append_zip_file):
      with tempfile.TemporaryDirectory() as temp_dir:
          zip_path = os.path.join(temp_dir, zip_file.name)
          output_zip_path = os.path.join(temp_dir, "output.zip")

          with open(zip_path, "wb") as f:
              f.write(zip_file.getbuffer())

          file_paths = []
          if files:
              for file in files:
                  file_path = os.path.join(temp_dir, file.name)
                  with open(file_path, "wb") as f:
                      f.write(file.getbuffer())
                  file_paths.append(file_path)

          if append_zip_file:
              append_zip_path = os.path.join(temp_dir, append_zip_file.name)
              with open(append_zip_path, "wb") as f:
                  f.write(append_zip_file.getbuffer())
              process_appended_zip(append_zip_path, output_zip_path)
          else:
              if file_type == "Both PDF and Word Document":
                  process_zip(zip_path, [fp for fp in file_paths if fp.endswith('.pdf')], output_zip_path, 'pdf')
                  process_zip(zip_path, [fp for fp in file_paths if fp.endswith('.docx')], output_zip_path, 'docx')
              else:
                  process_zip(zip_path, file_paths, output_zip_path, file_extension)
          
          with open(output_zip_path, "rb") as f:
              st.download_button("Download processed ZIP", f, file_name="processed_documents.zip")
          
          st.success(f"{file_type}s appended to documents in the output ZIP file successfully.")
  else:
      st.error("Please upload both a ZIP file and at least one file to append.")
